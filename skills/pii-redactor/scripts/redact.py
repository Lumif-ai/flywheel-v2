#!/usr/bin/env python3
"""
PII Redactor — Presidio-based document anonymization script.
Detects and redacts PII from text, PDF, and DOCX files.
Outputs redacted markdown + PII audit log.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import subprocess
import sys
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple


# ---------------------------------------------------------------------------
# Dependency management
# ---------------------------------------------------------------------------

REQUIRED_PACKAGES = {
    "presidio_analyzer": "presidio-analyzer",
    "presidio_anonymizer": "presidio-anonymizer",
    "spacy": "spacy",
}

OPTIONAL_PACKAGES = {
    "docx": "python-docx",
    "pdfplumber": "pdfplumber",
}


def ensure_installed(package_map: dict, label: str = "required"):
    """Install missing packages."""
    missing = []
    for module, pip_name in package_map.items():
        try:
            __import__(module)
        except ImportError:
            missing.append(pip_name)
    if missing:
        print(f"Installing {label} dependencies: {', '.join(missing)}")
        # Try --break-system-packages first (Python 3.11+), fall back to --user
        cmd = [sys.executable, "-m", "pip", "install", "-q"] + missing
        try:
            subprocess.check_call(cmd + ["--break-system-packages"])
        except subprocess.CalledProcessError:
            subprocess.check_call(cmd + ["--user"])


def ensure_spacy_model(model_name: str):
    """Download spaCy model if not already present."""
    import spacy

    try:
        spacy.load(model_name)
    except OSError:
        print(f"Downloading spaCy model: {model_name} (one-time setup)...")
        subprocess.check_call(
            [sys.executable, "-m", "spacy", "download", model_name]
        )


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------


def extract_text(file_path: str) -> str:
    """Extract text from PDF, DOCX, or plain text file."""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        ensure_installed({"pdfplumber": "pdfplumber"}, label="PDF")
        import pdfplumber

        pages = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
        return "\n\n".join(pages)

    elif ext in (".docx", ".doc"):
        ensure_installed({"docx": "python-docx"}, label="DOCX")
        import docx

        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n\n".join(paragraphs)

    else:
        # Plain text / markdown
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()


# ---------------------------------------------------------------------------
# PII Detection
# ---------------------------------------------------------------------------

DEFAULT_ENTITIES = [
    "PERSON",
    "EMAIL_ADDRESS",
    "PHONE_NUMBER",
    "US_SSN",
    "CREDIT_CARD",
    "IBAN_CODE",
    "US_PASSPORT",
    "US_DRIVER_LICENSE",
]

OPTIONAL_ENTITIES = [
    "LOCATION",
    "DATE_TIME",
    "ORGANIZATION",  # NER_WEAK included via ORGANIZATION
    "IP_ADDRESS",
    "URL",
    "MEDICAL_LICENSE",
    "US_BANK_NUMBER",
    "CRYPTO",
]

ALL_ENTITIES = DEFAULT_ENTITIES + OPTIONAL_ENTITIES


def detect_pii(text: str, entities: list[str], threshold: float, spacy_model: str):
    """Run Presidio Analyzer on text. Returns list of RecognizerResult."""
    from presidio_analyzer import AnalyzerEngine
    from presidio_analyzer.nlp_engine import NlpEngineProvider

    configuration = {
        "nlp_engine_name": "spacy",
        "models": [{"lang_code": "en", "model_name": spacy_model}],
    }
    provider = NlpEngineProvider(nlp_configuration=configuration)
    nlp_engine = provider.create_engine()

    analyzer = AnalyzerEngine(nlp_engine=nlp_engine, supported_languages=["en"])

    results = analyzer.analyze(
        text=text,
        entities=entities,
        language="en",
        score_threshold=threshold,
    )

    # Sort by position
    results = sorted(results, key=lambda r: r.start)
    return results


# ---------------------------------------------------------------------------
# PII Anonymization
# ---------------------------------------------------------------------------


def anonymize_text(
    text: str,
    results: list,
    mode: str = "replace",
    seed_mapping: dict[str, str] | None = None,
) -> tuple[str, list[dict]]:
    """
    Anonymize text using Presidio Anonymizer.
    Returns (anonymized_text, detailed_findings, reverse_map).

    In 'replace' mode, uses numbered tags: <PERSON_1>, <PERSON_2>, etc.

    If seed_mapping is provided (from a prior run's mapping.json), entities
    that match existing originals will reuse the same tags, ensuring
    consistency across documents (e.g., same person = same tag in both
    versions of a contract).
    """
    from presidio_anonymizer import AnonymizerEngine
    from presidio_anonymizer.entities import OperatorConfig

    engine = AnonymizerEngine()

    # Build consistent entity numbering
    entity_counters: dict[str, int] = defaultdict(int)
    entity_map: dict[str, str] = {}  # "PERSON:John Smith" -> "<PERSON_1>"
    reverse_map: dict[str, str] = {}  # "<PERSON_1>" -> "John Smith"

    # If seed mapping provided, pre-populate from prior run
    if seed_mapping:
        # seed_mapping is tag -> original, e.g. {"<PERSON_1>": "John Smith"}
        for tag, original in seed_mapping.items():
            # Parse entity type and number from tag like "<PERSON_1>"
            import re as _re
            m = _re.match(r"<([A-Z_]+)_(\d+)>", tag)
            if m:
                etype, num = m.group(1), int(m.group(2))
                key = f"{etype}:{original}"
                entity_map[key] = tag
                reverse_map[tag] = original
                entity_counters[etype] = max(entity_counters[etype], num)

    # Pre-assign numbers for consistency (same text = same tag)
    findings = []
    for result in results:
        original = text[result.start : result.end]
        key = f"{result.entity_type}:{original}"
        if key not in entity_map:
            entity_counters[result.entity_type] += 1
            n = entity_counters[result.entity_type]
            tag = f"<{result.entity_type}_{n}>"
            entity_map[key] = tag
            reverse_map[tag] = original

        # Calculate approximate line number
        line_num = text[: result.start].count("\n") + 1

        findings.append(
            {
                "entity_type": result.entity_type,
                "original": original,
                "replacement": entity_map[key],
                "confidence": round(result.score, 2),
                "start": result.start,
                "end": result.end,
                "line": line_num,
            }
        )

    if mode == "replace":
        # Manual replacement for consistent numbering (Presidio's built-in
        # replace doesn't support numbered tags natively)
        anonymized = text
        # Replace in reverse order to preserve positions
        for result in sorted(results, key=lambda r: r.start, reverse=True):
            original = anonymized[result.start : result.end]
            key = f"{result.entity_type}:{original}"
            anonymized = (
                anonymized[: result.start] + entity_map[key] + anonymized[result.end :]
            )
    elif mode == "hash":
        import hashlib

        anonymized = text
        for result in sorted(results, key=lambda r: r.start, reverse=True):
            original = anonymized[result.start : result.end]
            hashed = hashlib.sha256(original.encode()).hexdigest()[:12]
            anonymized = (
                anonymized[: result.start] + f"[{hashed}]" + anonymized[result.end :]
            )
            # Update finding
            for f in findings:
                if f["start"] == result.start:
                    f["replacement"] = f"[{hashed}]"
    elif mode == "mask":
        anonymized = text
        for result in sorted(results, key=lambda r: r.start, reverse=True):
            original = anonymized[result.start : result.end]
            if len(original) <= 2:
                masked = "*" * len(original)
            else:
                masked = original[0] + "*" * (len(original) - 2) + original[-1]
            anonymized = (
                anonymized[: result.start] + masked + anonymized[result.end :]
            )
            for f in findings:
                if f["start"] == result.start:
                    f["replacement"] = masked
    elif mode == "redact":
        anonymized = text
        for result in sorted(results, key=lambda r: r.start, reverse=True):
            anonymized = (
                anonymized[: result.start] + "[REDACTED]" + anonymized[result.end :]
            )
            for f in findings:
                if f["start"] == result.start:
                    f["replacement"] = "[REDACTED]"
    else:
        raise ValueError(f"Unknown mode: {mode}")

    return anonymized, findings, reverse_map


# ---------------------------------------------------------------------------
# Output generation
# ---------------------------------------------------------------------------


def generate_audit_log(
    source_file: str,
    findings: list[dict],
    threshold: float,
    spacy_model: str,
    mode: str,
    entities_used: list[str],
    entities_excluded: list[str],
) -> str:
    """Generate a markdown PII audit log."""
    date = datetime.now().strftime("%Y-%m-%d %H:%M")
    source_name = Path(source_file).name

    # Summary by type
    type_counts: dict[str, int] = defaultdict(int)
    for f in findings:
        type_counts[f["entity_type"]] += 1

    lines = [
        "# PII Audit Log",
        f"**Source:** {source_name}",
        f"**Date:** {date}",
        f"**Confidence threshold:** {threshold}",
        f"**spaCy model:** {spacy_model}",
        f"**Redaction mode:** {mode}",
        "",
        "## Summary",
        "",
        f"**Total PII entities found:** {len(findings)}",
        "",
        "| Entity Type | Count |",
        "|---|---|",
    ]
    for etype, count in sorted(type_counts.items()):
        lines.append(f"| {etype} | {count} |")

    lines += [
        "",
        "## Detailed Findings",
        "",
        "| # | Entity Type | Original | Replacement | Confidence | Line |",
        "|---|---|---|---|---|---|",
    ]
    for i, f in enumerate(findings, 1):
        orig = f["original"].replace("|", "\\|")
        repl = f["replacement"].replace("|", "\\|")
        lines.append(
            f"| {i} | {f['entity_type']} | `{orig}` | `{repl}` | {f['confidence']} | {f['line']} |"
        )

    lines += [
        "",
        "## Settings Used",
        f"- **Mode:** {mode}",
        f"- **Entities detected:** {', '.join(entities_used)}",
    ]
    if entities_excluded:
        lines.append(f"- **Entities excluded:** {', '.join(entities_excluded)}")

    lines.append("")
    lines.append("---")
    lines.append("*Generated by PII Redactor skill (Presidio-based)*")

    return "\n".join(lines)


def generate_reverse_map(reverse_map: dict[str, str]) -> str:
    """Generate a JSON reverse mapping file."""
    return json.dumps(reverse_map, indent=2, ensure_ascii=False)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main():
    parser = argparse.ArgumentParser(
        description="Detect and redact PII from documents using Microsoft Presidio."
    )
    parser.add_argument("input", help="Input file path (PDF, DOCX, TXT, MD)")
    parser.add_argument(
        "--threshold",
        type=float,
        default=0.7,
        help="Confidence threshold for PII detection (default: 0.7)",
    )
    parser.add_argument(
        "--output",
        help="Output file path (default: <input>_redacted.md)",
    )
    parser.add_argument(
        "--audit",
        help="Audit log file path (default: <input>_pii_audit.md)",
    )
    parser.add_argument(
        "--entities",
        help="Comma-separated entity types to detect (default: all default types)",
    )
    parser.add_argument(
        "--exclude-entities",
        help="Comma-separated entity types to skip",
    )
    parser.add_argument(
        "--mode",
        choices=["replace", "mask", "hash", "redact"],
        default="replace",
        help="Redaction mode (default: replace with type tags)",
    )
    parser.add_argument(
        "--spacy-model",
        default="en_core_web_lg",
        help="spaCy model to use (default: en_core_web_lg)",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Show what would be redacted without modifying anything",
    )
    parser.add_argument(
        "--save-mapping",
        action="store_true",
        help="Save reverse mapping file (<input>_mapping.json) for reconstruction",
    )
    parser.add_argument(
        "--seed-mapping",
        help="Path to existing mapping JSON from a prior redaction run. "
        "Ensures entity consistency across documents (same person = same tag).",
    )
    parser.add_argument(
        "--json",
        action="store_true",
        dest="json_output",
        help="Output detection results as JSON to stdout",
    )

    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    stem = input_path.stem
    parent = input_path.parent

    output_path = Path(args.output) if args.output else parent / f"{stem}_redacted.md"
    audit_path = Path(args.audit) if args.audit else parent / f"{stem}_pii_audit.md"
    mapping_path = parent / f"{stem}_mapping.json"

    # --- Install dependencies ---
    print("Checking dependencies...")
    ensure_installed(REQUIRED_PACKAGES)
    ensure_spacy_model(args.spacy_model)

    # --- Extract text ---
    print(f"Extracting text from: {input_path.name}")
    text = extract_text(str(input_path))
    if not text.strip():
        print("Error: No text extracted from file.", file=sys.stderr)
        sys.exit(1)
    print(f"Extracted {len(text):,} characters, {text.count(chr(10)) + 1} lines")

    # --- Determine entities ---
    entities_excluded = []
    if args.entities:
        entities = [e.strip().upper() for e in args.entities.split(",")]
    else:
        entities = list(DEFAULT_ENTITIES)
    if args.exclude_entities:
        entities_excluded = [e.strip().upper() for e in args.exclude_entities.split(",")]
        entities = [e for e in entities if e not in entities_excluded]

    # --- Detect PII ---
    print(f"Scanning for PII (threshold: {args.threshold}, entities: {len(entities)})...")
    results = detect_pii(text, entities, args.threshold, args.spacy_model)
    print(f"Found {len(results)} PII entities")

    if not results:
        print("\nNo PII detected. Document appears clean.")
        sys.exit(0)

    # --- Load seed mapping if provided ---
    seed_mapping = None
    if args.seed_mapping:
        seed_path = Path(args.seed_mapping)
        if not seed_path.exists():
            print(f"Warning: Seed mapping not found: {args.seed_mapping}", file=sys.stderr)
        else:
            with open(seed_path, "r", encoding="utf-8") as f:
                seed_mapping = json.load(f)
            print(f"Loaded seed mapping: {len(seed_mapping)} entities from {seed_path.name}")

    # --- Dry run: just show findings ---
    if args.dry_run or args.json_output:
        # Still need to build findings for display
        _, findings, reverse_map = anonymize_text(text, results, args.mode, seed_mapping=seed_mapping)

        if args.json_output:
            # Strip non-serializable fields
            print(json.dumps(findings, indent=2, ensure_ascii=False))
        else:
            # Pretty-print summary
            type_counts: dict[str, int] = defaultdict(int)
            type_samples: dict[str, list[str]] = defaultdict(list)
            for f in findings:
                type_counts[f["entity_type"]] += 1
                if len(type_samples[f["entity_type"]]) < 3:
                    type_samples[f["entity_type"]].append(f["original"])

            print("\n--- DRY RUN: PII Detection Summary ---\n")
            print(f"{'Entity Type':<25} {'Count':>5}   Samples")
            print("-" * 70)
            for etype in sorted(type_counts.keys()):
                samples = ", ".join(f'"{s}"' for s in type_samples[etype])
                print(f"{etype:<25} {type_counts[etype]:>5}   {samples}")
            print(f"\nTotal: {len(findings)} entities")
            print("\nRe-run without --dry-run to redact.")

        sys.exit(0)

    # --- Anonymize ---
    print(f"Redacting with mode: {args.mode}")
    anonymized, findings, reverse_map = anonymize_text(text, results, args.mode, seed_mapping=seed_mapping)

    # --- Write outputs ---
    # Redacted document
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(anonymized)
    print(f"Redacted document saved: {output_path}")

    # Audit log
    audit_content = generate_audit_log(
        source_file=str(input_path),
        findings=findings,
        threshold=args.threshold,
        spacy_model=args.spacy_model,
        mode=args.mode,
        entities_used=entities,
        entities_excluded=entities_excluded,
    )
    with open(audit_path, "w", encoding="utf-8") as f:
        f.write(audit_content)
    print(f"Audit log saved: {audit_path}")

    # Reverse mapping (opt-in)
    if args.save_mapping:
        with open(mapping_path, "w", encoding="utf-8") as f:
            f.write(generate_reverse_map(reverse_map))
        print(f"Reverse mapping saved: {mapping_path}")

    # --- Summary ---
    type_counts: dict[str, int] = defaultdict(int)
    for f in findings:
        type_counts[f["entity_type"]] += 1

    print("\n--- Redaction Summary ---")
    for etype in sorted(type_counts.keys()):
        print(f"  {etype}: {type_counts[etype]}")
    print(f"  Total: {len(findings)} entities redacted")


if __name__ == "__main__":
    main()
