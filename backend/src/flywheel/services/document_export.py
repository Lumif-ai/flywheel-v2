"""Document export service — PDF and DOCX generation from skill output.

Supports two paths:
- **Structured JSON** (one-pager, future skills): Builds branded documents
  from typed data using python-docx / WeasyPrint.
- **HTML fallback** (all other skills): Converts rendered_html to PDF via
  WeasyPrint, or strips to plain text for DOCX.

Public API:
    export_as_pdf(skill_name, output, rendered_html) -> bytes
    export_as_docx(skill_name, output, rendered_html) -> bytes
"""

from __future__ import annotations

import io
import json
import logging
from typing import Optional

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------


def export_as_pdf(
    skill_name: str,
    output: Optional[str],
    rendered_html: Optional[str],
) -> bytes:
    """Generate PDF bytes from skill output.

    Uses rendered_html (or generates standalone HTML) and converts via
    WeasyPrint. Falls back to a simple HTML wrapper if no rendered content.
    """
    html = rendered_html
    if not html and output:
        # Generate standalone HTML from raw output
        try:
            from flywheel.engines.output_renderer import render_output_standalone
            html = render_output_standalone(skill_name, output)
        except Exception:
            logger.warning("Standalone rendering failed for %s, using raw text", skill_name)
            html = _wrap_text_as_html(output, skill_name)

    if not html:
        html = _wrap_text_as_html("No content available.", skill_name)

    # Wrap fragment in full document if needed
    if "<html" not in html.lower():
        html = _wrap_fragment_as_document(html, skill_name)
    else:
        # For full documents, sanitize the body content in-place
        from flywheel.engines.output_renderer import sanitize_html
        import re as _re
        body_match = _re.search(r'(<body[^>]*>)(.*?)(</body>)', html, _re.DOTALL | _re.IGNORECASE)
        if body_match:
            sanitized_body = sanitize_html(body_match.group(2))
            html = html[:body_match.start(2)] + sanitized_body + html[body_match.end(2):]

    try:
        from weasyprint import HTML as WeasyprintHTML
        pdf_bytes = WeasyprintHTML(string=html).write_pdf()
        return pdf_bytes
    except ImportError:
        raise RuntimeError(
            "WeasyPrint is not installed. Install with: pip install weasyprint"
        )


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------


def export_as_docx(
    skill_name: str,
    output: Optional[str],
    rendered_html: Optional[str],
) -> bytes:
    """Generate DOCX bytes from skill output.

    For structured JSON (one-pager), builds a branded document with headings,
    tables, and styled content. For other skills, converts to plain text.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Brand colors
    coral = RGBColor(0xE9, 0x4D, 0x35)
    dark = RGBColor(0x12, 0x12, 0x12)
    muted = RGBColor(0x6B, 0x72, 0x80)
    light_gray = RGBColor(0x9C, 0xA3, 0xAF)

    # Try structured JSON path
    structured = _try_parse_structured(output)

    if structured and structured.get("document_type") == "value-prop-one-pager":
        _build_one_pager_docx(doc, structured, coral, dark, muted, light_gray)
    else:
        # Fallback: convert to plain text
        text = output or ""
        if not text and rendered_html:
            try:
                import html2text
                h = html2text.HTML2Text()
                h.ignore_links = False
                text = h.handle(rendered_html)
            except ImportError:
                text = rendered_html

        title = skill_name.replace("-", " ").replace("ctx-", "").title()
        heading = doc.add_heading(title, level=1)
        for run in heading.runs:
            run.font.color.rgb = dark

        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("## "):
                h = doc.add_heading(line[3:], level=2)
                for run in h.runs:
                    run.font.color.rgb = dark
            elif line.startswith("### "):
                h = doc.add_heading(line[4:], level=3)
                for run in h.runs:
                    run.font.color.rgb = dark
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# One-pager DOCX builder
# ---------------------------------------------------------------------------


def _build_one_pager_docx(doc, data, coral, dark, muted, light_gray):
    """Build a branded one-pager .docx from structured JSON data."""
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    # Set narrow margins for one-page fit
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Headline
    heading = doc.add_heading(data.get("headline", ""), level=1)
    for run in heading.runs:
        run.font.size = Pt(20)
        run.font.color.rgb = dark

    if data.get("subheadline"):
        p = doc.add_paragraph(data["subheadline"])
        p.runs[0].font.color.rgb = muted
        p.runs[0].font.size = Pt(11)

    # Stats banner as a table
    stats = data.get("stats_banner", [])
    if stats:
        table = doc.add_table(rows=2, cols=len(stats))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, stat in enumerate(stats):
            # Value row
            cell = table.cell(0, i)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stat.get("value", ""))
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = coral
            if stat.get("footnote_id") is not None:
                sup = p.add_run(str(stat["footnote_id"]))
                sup.font.size = Pt(8)
                sup.font.color.rgb = light_gray
                sup.font.superscript = True

            # Label row
            cell = table.cell(1, i)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stat.get("label", ""))
            run.font.size = Pt(8)
            run.font.color.rgb = light_gray

        # Remove table borders for clean look
        _remove_table_borders(table)
        doc.add_paragraph("")  # spacer

    # Problem columns
    columns = data.get("problem_columns", [])
    if columns:
        h = doc.add_heading("The Real Problem", level=2)
        for run in h.runs:
            run.font.color.rgb = coral
            run.font.size = Pt(10)

        table = doc.add_table(rows=1, cols=len(columns))
        for i, col in enumerate(columns):
            cell = table.cell(0, i)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(col.get("title", ""))
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = dark
            p2 = cell.add_paragraph(col.get("description", ""))
            for run in p2.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = muted

        _remove_table_borders(table)
        doc.add_paragraph("")

    # Outcomes
    outcomes = data.get("outcomes", [])
    if outcomes:
        h = doc.add_heading("What Changes", level=2)
        for run in h.runs:
            run.font.color.rgb = dark
            run.font.size = Pt(13)

        for outcome in outcomes:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(outcome.get("lead", ""))
            run.font.bold = True
            run.font.color.rgb = dark
            run.font.size = Pt(10)
            run2 = p.add_run(" " + outcome.get("detail", ""))
            run2.font.size = Pt(10)
            run2.font.color.rgb = muted

    # Comparison table
    comp = data.get("comparison_table", {})
    rows = comp.get("rows", [])
    col_names = comp.get("columns", ["Manual Process", "With Product"])
    if rows:
        h = doc.add_heading("Expected Outcomes", level=2)
        for run in h.runs:
            run.font.color.rgb = dark
            run.font.size = Pt(13)

        table = doc.add_table(rows=len(rows) + 1, cols=3)
        # Header
        headers = ["Metric"] + list(col_names)
        for i, header_text in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(header_text)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = coral if i == 2 else dark
            _shade_cell(cell, "FFF3F1" if i == 0 else "FFF3F1")

        for ri, row in enumerate(rows, 1):
            table.cell(ri, 0).text = row.get("metric", "")
            table.cell(ri, 1).text = row.get("manual", "")
            cell_product = table.cell(ri, 2)
            cell_product.text = ""
            p = cell_product.paragraphs[0]
            run = p.add_run(row.get("product", ""))
            run.font.color.rgb = coral
            run.font.bold = True
            run.font.size = Pt(9)

            for ci in range(3):
                for p in table.cell(ri, ci).paragraphs:
                    for run in p.runs:
                        if run.font.size is None:
                            run.font.size = Pt(9)

        doc.add_paragraph("")

    # Audit trail
    audit = data.get("audit_trail")
    if audit:
        p = doc.add_paragraph()
        run = p.add_run(audit.get("title", ""))
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = dark
        p2 = doc.add_paragraph(audit.get("description", ""))
        for run in p2.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = muted

    # Capability hint
    if data.get("capability_hint"):
        p = doc.add_paragraph(data["capability_hint"])
        for run in p.runs:
            run.font.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = light_gray

    # CTA
    cta = data.get("cta")
    if cta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cta.get("text", ""))
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = coral
        if cta.get("url"):
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p2.add_run(cta["url"])
            run.font.size = Pt(9)
            run.font.color.rgb = light_gray

    # Footnotes
    footnotes = data.get("footnotes", [])
    if footnotes:
        doc.add_paragraph("")
        for fn in footnotes:
            p = doc.add_paragraph()
            run = p.add_run(f"{fn.get('id', '')}. {fn.get('source', '')}")
            run.font.size = Pt(7)
            run.font.color.rgb = light_gray
            if fn.get("quality") and fn["quality"] != "first-party":
                run2 = p.add_run(f"  [{fn['quality']}]")
                run2.font.size = Pt(7)
                run2.font.color.rgb = light_gray


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _try_parse_structured(output: Optional[str]) -> Optional[dict]:
    """Try to parse output as structured JSON."""
    if not output or not output.strip().startswith("{"):
        return None
    try:
        data = json.loads(output)
        if isinstance(data, dict) and data.get("schema_version"):
            return data
    except (ValueError, TypeError):
        pass
    return None


def _wrap_text_as_html(text: str, skill_name: str) -> str:
    """Wrap plain text in a basic HTML document."""
    import html
    title = skill_name.replace("-", " ").replace("ctx-", "").title()
    escaped = html.escape(text).replace("\n", "<br>")
    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{title}</title>
<style>body {{ font-family: Inter, Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; color: #121212; font-size: 14px; line-height: 1.6; }}</style>
</head><body><h1>{title}</h1>{escaped}</body></html>"""


def _wrap_fragment_as_document(fragment: str, skill_name: str) -> str:
    """Wrap an HTML fragment in a full document with inline styles."""
    import html as html_mod
    from flywheel.engines.output_renderer import sanitize_html
    fragment = sanitize_html(fragment)  # XSS prevention -- LLM output is untrusted
    title = html_mod.escape(skill_name.replace("-", " ").replace("ctx-", "").title())
    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{title}</title>
<style>
body {{ font-family: Inter, -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; max-width: 960px; margin: 40px auto; padding: 20px; color: #121212; font-size: 14px; line-height: 1.6; }}
table {{ border-collapse: collapse; width: 100%; }}
th, td {{ padding: 8px 12px; border-bottom: 1px solid #E5E7EB; text-align: left; }}
th {{ font-weight: 600; }}
</style>
</head><body>{fragment}</body></html>"""


def _remove_table_borders(table):
    """Remove all borders from a docx table for a clean layout."""
    from docx.oxml.ns import qn

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "none",
            qn("w:sz"): "0",
            qn("w:space"): "0",
            qn("w:color"): "auto",
        })
        borders.append(el)
    tblPr.append(borders)


def _shade_cell(cell, color_hex: str):
    """Apply background shading to a docx table cell."""
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    tcPr.append(shading)
