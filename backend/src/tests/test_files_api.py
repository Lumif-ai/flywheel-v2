"""Tests for file upload, listing, and metadata endpoints.

Tests:
- TXT upload extracts text correctly
- File size limit (>10MB) returns 413
- Invalid mimetype returns 415
- List files returns paginated response without extracted_text
- Get file detail includes extracted_text
- PDF and DOCX extraction (unit tests on the service)
"""

from __future__ import annotations

import datetime
from io import BytesIO
from unittest.mock import AsyncMock, MagicMock, patch
from uuid import uuid4

import pytest
from fastapi.testclient import TestClient

from flywheel.api.deps import get_tenant_db, require_tenant
from flywheel.auth.jwt import TokenPayload
from flywheel.main import app
from flywheel.services.file_extraction import extract_text, validate_upload, MAX_FILE_SIZE

# ---------------------------------------------------------------------------
# Test constants
# ---------------------------------------------------------------------------

TEST_USER_ID = uuid4()
TEST_TENANT_ID = uuid4()
TEST_FILE_ID = uuid4()


def _make_user(tenant_id=TEST_TENANT_ID):
    return TokenPayload(
        sub=TEST_USER_ID,
        email="test@example.com",
        is_anonymous=False,
        app_metadata={"tenant_id": str(tenant_id), "role": "admin"},
    )


# ---------------------------------------------------------------------------
# Mock helpers
# ---------------------------------------------------------------------------


class MockResult:
    def __init__(self, value=None, values=None, scalar_val=None):
        self._value = value
        self._values = values or []
        self._scalar_val = scalar_val

    def scalar_one_or_none(self):
        return self._value

    def scalar(self):
        return self._scalar_val

    def scalars(self):
        return self

    def all(self):
        return self._values


class MockUploadedFile:
    def __init__(
        self,
        id=None,
        filename="test.txt",
        mimetype="text/plain",
        size_bytes=100,
        extracted_text="Hello world",
        storage_path="local://test",
        created_at=None,
    ):
        self.id = id or uuid4()
        self.filename = filename
        self.mimetype = mimetype
        self.size_bytes = size_bytes
        self.extracted_text = extracted_text
        self.storage_path = storage_path
        self.created_at = created_at or datetime.datetime(
            2026, 3, 20, tzinfo=datetime.timezone.utc
        )


def _mock_db(execute_side_effects=None):
    db = AsyncMock()
    if execute_side_effects:
        db.execute = AsyncMock(side_effect=execute_side_effects)
    else:
        db.execute = AsyncMock(return_value=MockResult())
    db.flush = AsyncMock()
    db.commit = AsyncMock()
    db.refresh = AsyncMock()
    db.add = MagicMock()
    return db


@pytest.fixture
def client():
    app.dependency_overrides = {}
    yield TestClient(app)
    app.dependency_overrides = {}


# ===========================================================================
# TestFileExtractionService (unit tests)
# ===========================================================================


class TestFileExtractionService:
    def test_validate_txt(self):
        """validate_upload accepts text/plain."""
        result = validate_upload("test.txt", "text/plain", 100)
        assert result == "text/plain"

    def test_validate_pdf(self):
        """validate_upload accepts application/pdf."""
        result = validate_upload("test.pdf", "application/pdf", 100)
        assert result == "application/pdf"

    def test_validate_docx(self):
        """validate_upload accepts DOCX mimetype."""
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        result = validate_upload("test.docx", mime, 100)
        assert result == mime

    def test_validate_rejects_invalid_mimetype(self):
        """validate_upload rejects unsupported mimetype."""
        with pytest.raises(ValueError, match="Unsupported"):
            validate_upload("test.zip", "application/zip", 100)

    def test_validate_rejects_oversized(self):
        """validate_upload rejects files over 10MB."""
        with pytest.raises(ValueError, match="too large"):
            validate_upload("big.txt", "text/plain", MAX_FILE_SIZE + 1)

    @pytest.mark.asyncio
    async def test_extract_txt(self):
        """extract_text handles TXT files."""
        content = b"Hello, world!"
        result = await extract_text(content, "text/plain")
        assert result == "Hello, world!"

    @pytest.mark.asyncio
    async def test_extract_txt_with_unicode(self):
        """extract_text handles UTF-8 TXT with special characters."""
        content = "Hello \u00e9\u00e0\u00fc".encode("utf-8")
        result = await extract_text(content, "text/plain")
        assert "\u00e9" in result

    @pytest.mark.asyncio
    async def test_extract_unsupported_mimetype(self):
        """extract_text raises ValueError for unsupported types."""
        with pytest.raises(ValueError, match="Unsupported"):
            await extract_text(b"data", "application/zip")

    @pytest.mark.asyncio
    async def test_extract_docx(self):
        """extract_text handles DOCX files."""
        import docx

        doc = docx.Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        buf = BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        result = await extract_text(content, mime)
        assert "First paragraph" in result
        assert "Second paragraph" in result


# ===========================================================================
# TestUploadEndpoint
# ===========================================================================


class TestUploadEndpoint:
    def test_upload_txt_file(self, client):
        """POST /files/upload with TXT returns 201 with extracted text length."""
        user = _make_user()
        mock_db = _mock_db()

        def mock_refresh(obj):
            obj.id = TEST_FILE_ID
            obj.created_at = datetime.datetime(2026, 3, 20, tzinfo=datetime.timezone.utc)

        mock_db.refresh = AsyncMock(side_effect=mock_refresh)

        app.dependency_overrides[require_tenant] = lambda: user
        app.dependency_overrides[get_tenant_db] = lambda: mock_db

        resp = client.post(
            "/api/v1/files/upload",
            files={"file": ("test.txt", b"Hello from test file", "text/plain")},
        )
        assert resp.status_code == 201
        data = resp.json()
        assert data["filename"] == "test.txt"
        assert data["mimetype"] == "text/plain"
        assert data["size_bytes"] == 20
        assert data["extracted_text_length"] == 20

    def test_upload_oversized_returns_413(self, client):
        """POST /files/upload with >10MB returns 413."""
        user = _make_user()
        mock_db = _mock_db()

        app.dependency_overrides[require_tenant] = lambda: user
        app.dependency_overrides[get_tenant_db] = lambda: mock_db

        # Create content just over the limit
        content = b"x" * (MAX_FILE_SIZE + 1)
        resp = client.post(
            "/api/v1/files/upload",
            files={"file": ("big.txt", content, "text/plain")},
        )
        assert resp.status_code == 413

    def test_upload_invalid_mimetype_returns_415(self, client):
        """POST /files/upload with .zip returns 415."""
        user = _make_user()
        mock_db = _mock_db()

        app.dependency_overrides[require_tenant] = lambda: user
        app.dependency_overrides[get_tenant_db] = lambda: mock_db

        resp = client.post(
            "/api/v1/files/upload",
            files={"file": ("test.zip", b"PK\x03\x04fake", "application/zip")},
        )
        assert resp.status_code == 415

    def test_upload_requires_auth(self, client):
        """POST /files/upload without auth returns 401."""
        resp = client.post(
            "/api/v1/files/upload",
            files={"file": ("test.txt", b"hello", "text/plain")},
        )
        assert resp.status_code == 401


# ===========================================================================
# TestListFiles
# ===========================================================================


class TestListFiles:
    def test_list_files_paginated(self, client):
        """GET /files/ returns paginated list without extracted_text."""
        user = _make_user()
        files = [MockUploadedFile() for _ in range(2)]
        mock_db = _mock_db([
            MockResult(scalar_val=2),
            MockResult(values=files),
        ])

        app.dependency_overrides[require_tenant] = lambda: user
        app.dependency_overrides[get_tenant_db] = lambda: mock_db

        resp = client.get("/api/v1/files/")
        assert resp.status_code == 200
        data = resp.json()
        assert data["total"] == 2
        assert len(data["items"]) == 2
        # List should NOT include extracted_text
        assert "extracted_text" not in data["items"][0]

    def test_list_files_empty(self, client):
        """GET /files/ returns empty list when no files."""
        user = _make_user()
        mock_db = _mock_db([
            MockResult(scalar_val=0),
            MockResult(values=[]),
        ])

        app.dependency_overrides[require_tenant] = lambda: user
        app.dependency_overrides[get_tenant_db] = lambda: mock_db

        resp = client.get("/api/v1/files/")
        assert resp.status_code == 200
        data = resp.json()
        assert data["total"] == 0
        assert data["items"] == []


# ===========================================================================
# TestGetFile
# ===========================================================================


class TestGetFile:
    def test_get_file_detail(self, client):
        """GET /files/{id} returns file with extracted_text."""
        user = _make_user()
        f = MockUploadedFile(id=TEST_FILE_ID, extracted_text="Full document text here")
        mock_db = _mock_db([MockResult(value=f)])

        app.dependency_overrides[require_tenant] = lambda: user
        app.dependency_overrides[get_tenant_db] = lambda: mock_db

        resp = client.get(f"/api/v1/files/{TEST_FILE_ID}")
        assert resp.status_code == 200
        data = resp.json()
        assert data["id"] == str(TEST_FILE_ID)
        assert data["extracted_text"] == "Full document text here"
        assert data["extracted_text_length"] == len("Full document text here")

    def test_get_file_not_found(self, client):
        """GET /files/{id} returns 404 for missing file."""
        user = _make_user()
        mock_db = _mock_db([MockResult(value=None)])

        app.dependency_overrides[require_tenant] = lambda: user
        app.dependency_overrides[get_tenant_db] = lambda: mock_db

        resp = client.get(f"/api/v1/files/{uuid4()}")
        assert resp.status_code == 404
